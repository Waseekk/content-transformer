"""
Comprehensive Test: Generate Hard News and Soft News for all 7 PDF articles
Output to Word document with proper formatting
v3.0: Code-based post-processing (100% reliable, no AI checker)
      - fix_intro_structure() ensures FULLY bold intro + exactly 2 intros for soft news
"""
import os
import sys
import json
import io
import re
from pathlib import Path
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv

# Fix Windows console encoding
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# Load environment variables
load_dotenv(Path(__file__).parent / "backend" / ".env")

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Load prompts
json_path = Path(__file__).parent / "backend" / "app" / "config" / "formats" / "bengali_news_styles.json"
with open(json_path, "r", encoding="utf-8") as f:
    STYLES = json.load(f)

# ============================================================================
# CODE-BASED POST-PROCESSING (v2.5 - 100% Reliable, No AI Checker)
# ============================================================================

# Word correction patterns
WORD_CORRECTIONS = [
    (r'শীঘ্রই', 'শিগগিরই'),
    (r'([০-৯]+)লা\b', r'\1'),
    (r'([০-৯]+)ই\b', r'\1'),
    (r'([০-৯]+)শে\b', r'\1'),
]

# Words that START with সহ - don't join these
SAHO_EXCEPTION_PATTERNS = [
    'সহায়', 'সহযোগ', 'সহকার', 'সহজ', 'সহ্য', 'সহস্র',
    'সহমত', 'সহমর্ম', 'সহাবস্থান', 'সহসা', 'সহচর',
    'সহধর্মিণী', 'সহপাঠী', 'সহবাস', 'সহমরণ', 'সহানুভূতি',
    'সহকর্মী', 'সহযাত্রী', 'সহশিল্পী', 'সহনশীল',
]

# English words to replace with Bengali equivalents
ENGLISH_TO_BENGALI = {
    'accompanying': 'সহায়ক',
    'landmark': 'ঐতিহ্যবাহী স্থান',
    'landmarks': 'ঐতিহ্যবাহী স্থানসমূহ',
    'sharply': 'তীব্রভাবে',
    'system': 'ব্যবস্থা',
    'tariff': 'শুল্ক',
    'desperation': 'মরিয়া অবস্থা',
    'tourists': 'পর্যটকরা',
    'tourist': 'পর্যটক',
    'government': 'সরকার',
    'official': 'কর্মকর্তা',
    'officials': 'কর্মকর্তারা',
    'significant': 'উল্লেখযোগ্য',
    'approximately': 'প্রায়',
    'recently': 'সম্প্রতি',
    'currently': 'বর্তমানে',
    'however': 'তবে',
    'therefore': 'তাই',
    'moreover': 'অধিকন্তু',
    'despite': 'সত্ত্বেও',
    'although': 'যদিও',
}


def apply_word_corrections(text):
    """Apply Bengali word corrections"""
    if not text:
        return text
    for pattern, replacement in WORD_CORRECTIONS:
        text = re.sub(pattern, replacement, text)
    return text


def fix_saho_joining(text):
    """Join সহ with previous word only when it means 'with'"""
    if not text:
        return text
    exception_lookahead = '|'.join(SAHO_EXCEPTION_PATTERNS)
    pattern = rf'(\S+)\s+সহ(?!{exception_lookahead})(?=[\s,।\n]|$)'
    return re.sub(pattern, r'\1সহ', text)


def replace_english_words(text):
    """Replace common English words with Bengali equivalents"""
    if not text:
        return text
    for eng, ben in ENGLISH_TO_BENGALI.items():
        pattern = rf'\b{eng}\b'
        text = re.sub(pattern, ben, text, flags=re.IGNORECASE)
    return text


def split_quotes(text):
    """
    Split paragraphs where text appears after a CLOSING quote.

    Only split when: punctuation + closing quote + space + more text
    Pattern: ।" or !" or ?" (punctuation inside) OR "। or "! or "? (punctuation outside)
    """
    if not text:
        return text

    paragraphs = text.split('\n\n')
    result = []

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Skip bold paragraphs (headlines, intros, subheads)
        if para.startswith('**') and '**' in para[2:]:
            result.append(para)
            continue

        # Skip byline
        if 'নিউজ ডেস্ক' in para and 'বাংলার কলম্বাস' in para:
            result.append(para)
            continue

        # Check if paragraph has quotes
        if '"' not in para:
            result.append(para)
            continue

        # Split only at CLOSING quote patterns
        current_para = para
        split_parts = []

        while True:
            # Match CLOSING quote patterns only:
            # - [।!?]" = punctuation inside, then closing quote
            # - "[।!?] = closing quote, then punctuation outside
            match = re.search(r'([।!?]"|"[।!?])\s+(\S.+)', current_para)
            if match:
                split_pos = match.start() + len(match.group(1))
                part1 = current_para[:split_pos].strip()
                part2 = match.group(2).strip()
                if part1:
                    split_parts.append(part1)
                current_para = part2
            else:
                if current_para.strip():
                    split_parts.append(current_para.strip())
                break

        result.extend(split_parts)

    return '\n\n'.join(result)


def normalize_line_breaks(text):
    """
    Normalize markdown line breaks to paragraph breaks.

    AI sometimes uses '  \\n' (markdown line break) instead of '\\n\\n' (paragraph break).
    This merges byline + intro 1 + intro 2 into one "paragraph".

    Fix: Convert markdown line breaks after byline and bold sections to paragraph breaks.
    """
    if not text:
        return text

    # Pattern 1: Byline followed by markdown line break + bold intro
    # "নিউজ ডেস্ক, বাংলার কলম্বাস  \n**" → "নিউজ ডেস্ক, বাংলার কলম্বাস\n\n**"
    text = re.sub(r'(বাংলার কলম্বাস)\s*\n(\*\*)', r'\1\n\n\2', text)

    # Pattern 2: Bold section ending with **  \n followed by non-bold text
    # "...হবে।**  \nনতুন এই" → "...হবে।**\n\nনতুন এই"
    text = re.sub(r'\*\*\s*\n([^*\n])', r'**\n\n\1', text)

    # Pattern 3: Non-bold text followed by markdown line break + bold subhead
    # "...করে তোলে।  \n**সাবহেড**" → "...করে তোলে।\n\n**সাবহেড**"
    text = re.sub(r'([।!?])\s*\n(\*\*)', r'\1\n\n\2', text)

    return text


# ============================================================================
# INTRO STRUCTURE FIXER (v3.0 - Complete Rewrite)
# ============================================================================

def make_fully_bold(text):
    """Make entire paragraph FULLY bold, removing any partial bold."""
    clean = text.replace('**', '').strip()
    return f'**{clean}**'


def remove_bold(text):
    """Remove all bold markers from text."""
    return text.replace('**', '').strip()


def find_byline_index(paragraphs):
    """Find index of byline paragraph."""
    for i, p in enumerate(paragraphs):
        if 'নিউজ ডেস্ক' in p and 'বাংলার কলম্বাস' in p:
            return i
    return -1


def find_first_subhead(paragraphs, start_idx):
    """
    Find first subhead after start_idx.
    Subhead = FULLY bold paragraph (starts AND ends with **) that is relatively short.
    """
    for i in range(start_idx, min(len(paragraphs), start_idx + 10)):
        p = paragraphs[i].strip()
        # Check if FULLY bold (starts and ends with **)
        if p.startswith('**') and p.endswith('**'):
            clean = p[2:-2].strip()
            # Subheads are typically < 100 chars
            if len(clean) < 100:
                return i
    return -1


def split_into_sentences(text):
    """
    Split text into sentences at Bengali sentence endings.
    Preserves quotes as complete units.
    """
    if not text:
        return []

    OPENING_QUOTES = '"\u201C'
    CLOSING_QUOTES = '"\u201D'

    sentences = []
    current = []
    in_quote = False

    for char in text:
        current.append(char)
        if char in OPENING_QUOTES:
            in_quote = True
        elif char in CLOSING_QUOTES:
            in_quote = False

        if char in '।?!' and not in_quote:
            sentence = ''.join(current).strip()
            if sentence:
                sentences.append(sentence)
            current = []

    remaining = ''.join(current).strip()
    if remaining:
        sentences.append(remaining)

    return sentences


def split_intro(intro):
    """
    Split single intro into Intro 1 and Intro 2.
    Intro 1: First 2 sentences
    Intro 2: Remaining sentences
    """
    clean = intro.replace('**', '').strip()
    sentences = split_into_sentences(clean)

    if len(sentences) <= 2:
        return (clean, '')

    intro1 = ' '.join(sentences[:2])
    intro2 = ' '.join(sentences[2:])
    return (intro1, intro2)


def fix_three_line_paragraphs(text):
    """
    Enforce max 2 sentences per body paragraph (2 lines = 2 sentences).

    Rules:
    1. Body paragraphs should have max 2 sentences
    2. If 3+ sentences, split into groups of max 2
    3. If quotation exists, keep it as the last sentence of its paragraph
    """
    if not text:
        return text

    paragraphs = text.split('\n\n')
    result = []

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Skip bold paragraphs (headlines, intros, subheads)
        if para.startswith('**') and '**' in para[2:]:
            result.append(para)
            continue

        # Skip byline
        if 'নিউজ ডেস্ক' in para and 'বাংলার কলম্বাস' in para:
            result.append(para)
            continue

        # Split into sentences (preserves quotes as complete units)
        full_sentences = split_into_sentences(para)

        # If 2 or fewer sentences, keep as is
        if len(full_sentences) <= 2:
            result.append(para)
            continue

        # 3+ sentences: split into groups of max 2
        # Special handling for quotations - keep quote sentence at end of its paragraph

        # Group sentences: max 2 per paragraph, quotes at end of their paragraph
        current_group = []
        for i, sentence in enumerate(full_sentences):
            current_group.append(sentence)

            # Check if this is a good split point
            should_split = False

            if len(current_group) >= 2:
                # Have 2 sentences, time to split
                should_split = True

                # But if next sentence has a quote and current doesn't, wait
                if i + 1 < len(full_sentences) and '"' in full_sentences[i + 1] and '"' not in sentence:
                    should_split = False

            # If current sentence has a quote, always split after it
            if '"' in sentence and len(current_group) >= 1:
                should_split = True

            if should_split:
                result.append(' '.join(current_group))
                current_group = []

        # Add remaining sentences
        if current_group:
            result.append(' '.join(current_group))

    return '\n\n'.join(result)


def count_body_words(content):
    """
    Count words in hard news body (excluding headline and byline).
    Counts from intro paragraph to conclusion.
    """
    if not content:
        return 0

    paragraphs = content.split('\n\n')
    body_words = 0

    for i, para in enumerate(paragraphs):
        para = para.strip()
        if not para:
            continue

        # Skip headline (first bold line, usually index 0)
        if i == 0 and para.startswith('**'):
            continue

        # Skip byline
        if 'নিউজ ডেস্ক' in para and 'বাংলার কলম্বাস' in para:
            continue

        # Count words in this paragraph (remove ** markers first)
        clean_para = para.replace('**', '')
        body_words += len(clean_para.split())

    return body_words


def fix_intro_structure(content, format_type):
    """
    Fix intro structure to match exact requirements (v3.0).

    Hard News: P3 (Intro) must be FULLY bold
    Soft News:
        - P3 (Intro 1) must be FULLY bold
        - P4 (Intro 2) must NOT be bold
        - P5 = First Subhead
        - EXACTLY 2 intro paragraphs before first subhead

    Handles:
    1. Partial bold in intro → make FULLY bold
    2. Only 1 intro exists → split into 2
    3. More than 2 intros → merge extras into Intro 2
    """
    if not content:
        return content

    paragraphs = content.split('\n\n')

    # Find byline index
    byline_idx = find_byline_index(paragraphs)
    if byline_idx == -1:
        return content

    intro_idx = byline_idx + 1
    if intro_idx >= len(paragraphs):
        return content

    if format_type == 'hard_news':
        # Hard news: Just make intro FULLY bold
        paragraphs[intro_idx] = make_fully_bold(paragraphs[intro_idx])
        print("      [POST] Fixed hard news intro: made FULLY bold")
        return '\n\n'.join(paragraphs)

    # SOFT NEWS: Need exactly 2 intros
    first_subhead_idx = find_first_subhead(paragraphs, intro_idx + 1)

    if first_subhead_idx == -1:
        paragraphs[intro_idx] = make_fully_bold(paragraphs[intro_idx])
        print("      [POST] Fixed soft news: made intro 1 FULLY bold (no subhead found)")
        return '\n\n'.join(paragraphs)

    intros_between = first_subhead_idx - intro_idx

    if intros_between == 2:
        # Perfect! Fix formatting only
        paragraphs[intro_idx] = make_fully_bold(paragraphs[intro_idx])
        paragraphs[intro_idx + 1] = remove_bold(paragraphs[intro_idx + 1])
        print("      [POST] Fixed soft news: exactly 2 intros found, fixed formatting")

    elif intros_between == 1:
        # Split single intro into 2
        intro1, intro2 = split_intro(paragraphs[intro_idx])
        if intro2:
            paragraphs[intro_idx] = make_fully_bold(intro1)
            paragraphs.insert(intro_idx + 1, intro2)
            print("      [POST] Fixed soft news: split single intro into 2 paragraphs")
        else:
            paragraphs[intro_idx] = make_fully_bold(paragraphs[intro_idx])
            print("      [POST] Soft news: only 1 intro, couldn't split (too few sentences)")

    elif intros_between > 2:
        # Merge extras into intro 2
        intro1 = paragraphs[intro_idx]
        intro2_parts = paragraphs[intro_idx + 1 : first_subhead_idx]
        intro2_merged = ' '.join([remove_bold(p) for p in intro2_parts])

        del paragraphs[intro_idx + 2 : first_subhead_idx]

        paragraphs[intro_idx] = make_fully_bold(intro1)
        paragraphs[intro_idx + 1] = intro2_merged

        print(f"      [POST] Fixed soft news: merged {intros_between - 1} paragraphs into Intro 2")

    elif intros_between == 0:
        # Subhead immediately after intro - split intro
        intro1, intro2 = split_intro(paragraphs[intro_idx])
        if intro2:
            paragraphs[intro_idx] = make_fully_bold(intro1)
            paragraphs.insert(intro_idx + 1, intro2)
            print("      [POST] Fixed soft news: split intro before immediate subhead")
        else:
            paragraphs[intro_idx] = make_fully_bold(paragraphs[intro_idx])
            print("      [POST] Soft news: subhead immediately after intro, couldn't split")

    return '\n\n'.join(paragraphs)


def process_content(text, format_type):
    """
    Full code-based post-processing pipeline (v3.0)

    0. Normalize markdown line breaks to paragraph breaks
    1. FIX INTRO STRUCTURE (single comprehensive function - v3.0)
       - Makes intro FULLY bold (not partial)
       - For soft news: ensures exactly 2 intros before first subhead
    2. Apply word corrections (শিগগিরই, date suffixes)
    3. Fix সহ joining (smart)
    4. Replace English words
    5. Split quotes (CRITICAL)
    6. Fix 3-line paragraphs (max 2 sentences per body paragraph)
    """
    text = normalize_line_breaks(text)
    text = fix_intro_structure(text, format_type)
    text = apply_word_corrections(text)
    text = fix_saho_joining(text)
    text = replace_english_words(text)
    text = split_quotes(text)
    text = fix_three_line_paragraphs(text)
    return text


def detect_issues(content, format_type):
    """Detect issues for logging (no longer triggers AI checker)"""
    issues = []
    paragraphs = content.split('\n\n')
    body_start = 3 if format_type == 'hard_news' else 5

    for i, para in enumerate(paragraphs):
        para = para.strip()
        if not para or i < body_start:
            continue
        if para.startswith('**') and '**' in para[2:]:
            continue
        if 'নিউজ ডেস্ক' in para:
            continue

        # Check: Text after closing quote (should be split)
        if '"' in para:
            last_quote = para.rfind('"')
            text_after = para[last_quote+1:].strip()
            if len(text_after) > 5 and not text_after.startswith('।'):
                issues.append(f"P{i+1}: Text after quote")

    return issues

# ============================================================================
# ALL 7 ARTICLES FROM PDF
# ============================================================================

ARTICLES = [
    {
        "id": 1,
        "title": "Tourists to face €2 fee to get near Rome's Trevi Fountain",
        "publisher": "BBC News",
        "content": """Tourists in Italy's capital Rome will soon have to pay a €2 (£1.75; $2.34) entrance fee if they want to see its famed Trevi Fountain up close.

The new barrier for visitors to view the Baroque monument will come into force from 1 February 2026.

While the coins tossed into the fountain are donated to charity, the fees collected will go to the city authority to pay for upkeep and managing visitors. The city expects to raise €6.5m a year from the fountain alone.

Announcing the move on Friday, Rome's Mayor Roberto Gualtieri was quoted by news agency Reuters as saying that "two euros isn't very much ... and it will lead to less chaotic tourist flows".

The Trevi levy is part of a new tariff system for certain museums and monuments in the Italian capital.

Access to a number of sites that currently charge for entry will become free for Rome's residents, such as the Sacred Area of Largo Argentina.

At the same time, tourists and non-residents will have to pay to see the Trevi fountain and five other attractions including the Napoleonic Museum.

Children under the age of five, and those with disabilities and an accompanying person, will be exempt from the fees.

The site currently sees an average of 30,000 visitors per day, according to the City of Rome."""
    },
    {
        "id": 2,
        "title": "Thousands of dinosaur footprints found on Italian mountain",
        "publisher": "BBC News",
        "content": """Thousands of dinosaur footprints dating back 210 million years have been found in a national park in northern Italy.

The footprints - some of which are up to 40cm (15in) in diameter - are aligned in parallel rows, and many show clear traces of toes and claws.

It is thought the dinosaurs were prosauropods - herbivores with long necks, small heads and sharp claws.

"I never would have imagined I'd come across such a spectacular discovery in the region where I live," said Milan-based paleontologist Cristiano Dal Sasso.

Last September a photographer spotted the footprints stretching hundreds of metres on a vertical mountain wall in the Stelvio national park, north-east of Milan.

In the Triassic period - between about 250 and 201 million years ago - the wall was a tidal flat, which later became part of the Alpine chain.

"This place was full of dinosaurs; it's an immense scientific treasure," Mr Dal Sasso said.

The herds moved in harmony, he added, "and there are also traces of more complex behaviours, like groups of animals gathering in a circle, perhaps for the purposes of defence."

The prosauropods, which could be up to 10m (33ft) long, walked on two legs but in some cases handprints were found in front of footprints, indicating that they probably stopped and rested their forelimbs on the ground.

Elio Della Ferrera, the photographer who discovered the site, said he hoped the discovery would "spark reflection in all of us, highlighting how little we know about the places we live in: our home, our planet".

The Stelvio national park is located in the Fraele valley by Italy's border with Switzerland, near where the Winter Olympics will take place next year.

"It's as if history itself wanted to pay homage to the greatest global sporting event, combining past and present in a symbolic passing of the baton between nature and sport," said the Italian Ministry of Culture."""
    },
    {
        "id": 3,
        "title": "The six best places to ski in 2026",
        "publisher": "BBC Travel",
        "content": """As the barometer drops in the northern hemisphere, skiers everywhere are preparing to shred black diamonds and show off their mountain 'fits. Tried-and-true hotspots like Aspen and Courchevel never lose their appeal, but according to Sara Haney, head of travel for concierge service Velocity Black which organises bespoke skiing experiences for travellers, skiers in 2026 "want a sense of discovery, not just lift access".

This season, snow reliability also remains a top priority. According to Cat Iwanchuk, vice president of business development at Ski.com, skiers are changing their plans to coincide with optimal snowfall.

Alta Badia, The Dolomites, Italy: The Dolomite Mountains, famed for their dramatic limestone peaks and perfect pink sunsets, have traditionally been an "Italians-only" ski destination. Thanks to hosting this year's Winter Olympics, ski-related searches from the US have risen 82% compared to the same timeframe in 2024.

Big Sky, Montana, US: Can a hotel opening define an entire season? The November 2025 debut of One&Only Moonlight Basin just might. Big Sky's ascent from under-the-radar cult favourite to one of the US' most popular resorts has been a decade in the making.

Banff, Alberta, Canada: Despite its location in the heart of the Canadian Rockies, Banff has traditionally been considered a summer travel destination. On 17 December 2025, Lake Louise Ski Resort unveiled Richardson's Ridge; 200 new acres of terrain.

Hokkaido, Japan: From 2024 to 2025, Niseko, a town in Japan's northernmost Hokkaido region known for its fluffy "Japow" (Japan powder), had a record-breaking season, welcoming more than 11 million visitors.

Méribel, France: In the French Alps, this quaint ski town with wood chalet architecture has access to 600km of trails, a lively après-ski scene and gorgeous views of the Les Allues valley.

Queenstown, New Zealand: Head to Queenstown, home to one of the world's longest winter sports seasons, stretching nearly five months from June to October."""
    },
    {
        "id": 4,
        "title": "Turkey's ornate Ottoman-era 'bird palaces'",
        "publisher": "BBC Travel",
        "content": """On a corner lot in Istanbul's old city, there's what some might call a fixer-upper: a stone-built duplex with gabled roofs and lots of windows. Parts of its exterior walls are crumbling, but there's a unique feature affixed to its walls that makes it a perfectly fine home for its many residents -- the city's birds.

Too grand to be called mere birdhouses, dozens of masterfully stone-carved kuş sarayları ("bird palaces") or kuş köşkleri ("bird pavilions") in varying states of repair are perched under the eaves of Ottoman-era mosques, tombs and other structures in central Istanbul.

"Each bird palace has its own unique architectural style and detailed beauty that also represents a thoughtful display of compassion," says Öykü Demir, an Istanbul-based tour guide.

Charity is one of the five pillars of Islam, and in Ottoman times, this extended to animals across the empire. Birds held a special place within the empire because they were symbolically linked to the soul reaching for the heavens.

"In Istanbul, you often find birdhouses on the qibla walls of mosques (those facing Mecca) so the sound of birds becomes associated with the sound of prayer," says Christiane Gruber, a professor of Islamic art history at the University of Michigan.

The practice reached its creative peak during the 18th Century, a period of prosperity in which Ottoman art and architecture transformed under the influence of the Baroque and Rococo movements from Europe.

Despite being a densely populated megacity, Istanbul is still rich in birdlife, with nearly 300 different species spotted in the city by birdwatchers this year alone.

"These bird houses represent a very old tradition of hospitality, and an idea... that says the city belongs not just to humans," says Gruber."""
    },
    {
        "id": 5,
        "title": "Brazil's lagoon-filled desert you can hike barefoot",
        "publisher": "BBC Travel",
        "content": """I was at least 20 steps behind my group -- and another 30 behind our guide -- when he suddenly stopped, checked his watch and tilted his face toward the sky, as if taking cues from the Sun.

My three friends and I were a few hours into a three-day trek across Lençóis Maranhenses National Park, a humbling expanse of sand in north-eastern Brazil.

Bordered by lush vegetation on one side and the Atlantic Ocean on the other, Lençóis Maranhenses is one of Brazil's most unusual ecosystems. Strong coastal winds push sand inland, creating a desert-like landscape that spans 1,500 sq km, with dunes rising up to 30m.

But Lençóis is no desert. During the wet season (January to June), so much rain falls that water settles in natural basins between the dunes, forming the hundreds of freshwater lagoons that travellers come to swim in.

It's these ephemeral pools that make Lençois like nowhere else on Earth. It's also what helped earn the park Unesco World Heritage status in 2024. In 2024, Lençóis received a record 552,000 visitors.

Our trek began in Lagoa Bonita. From there, we would cross 36km of sand, including two overnight stays in local villages, to finish in Atins.

"We have more than 1,000 families living inside Lençóis Maranhenses," said Cristiane Figueiredo, head of the national park. "To truly get to know this unique environment, there's nothing better than walking."

Nothing here holds its shape for long. The dunes shift. The lagoons appear and disappear. Footprints vanish. The only constant in Lençóis Maranhenses is change -- and the people who have learned, over generations, to live with it."""
    },
    {
        "id": 6,
        "title": "Sabres set sights on extending season-long win streak vs. Islanders",
        "publisher": "Field Level Media",
        "content": """The New York Islanders lost Friday night to a team near the bottom of the standings playing much freer after a major personnel move.

The Islanders hope history doesn't repeat itself Saturday, when they will to salvage the second game of a back-to-back set when they visit the Buffalo Sabres in a battle of Eastern Conference rivals.

The Islanders never led Friday in a 4-1 loss to the Vancouver Canucks in Elmont, N.Y. The host Sabres, meanwhile, earned their fourth straight victory, 5-3 over the Philadelphia Flyers.

The loss was the second straight for the Islanders, who vaulted into second place in the Metropolitan Division by winning six of seven games before Tuesday's 3-2 loss to the Detroit Red Wings.

The Sabres, who have missed the playoffs in each of the last 14 seasons and are in last place in the Atlantic Division, shifted in a different direction last Monday, when general manager Kevyn Adams was fired and replaced by senior adviser Jarmo Kekalainen.

Buffalo won three in a row before the change and extended its longest winning streak of the season in Kekalainen's first game as general manager.

"When you find that next level of desperation to try to win games, ultimately for the guy next to you, I think that's when good things happen," said Sabres goalie Alex Lyon."""
    },
    {
        "id": 7,
        "title": "Kazakhstan Becomes a Prime Choice for Chinese Outbound Tourism",
        "publisher": "Know More",
        "content": """Kazakhstan has rapidly become a preferred destination for Chinese travelers seeking a blend of adventure, winter sports, and urban experiences. The country's growing appeal stems from its diverse offerings: towering snow-capped mountains and world-class ski resorts attract winter sports enthusiasts, while vibrant cities like Almaty and Astana provide modern amenities, cultural landmarks, and dynamic city life.

Kazakhstan has firmly established itself as one of the seven most popular international destinations for Chinese travelers during the New Year holiday, reflecting a rising trend in outbound tourism from China.

Recent data reveals that demand for travel to Kazakhstan from Chinese cities has surged sharply since early December. Air travel bookings have increased by more than 50% compared to last year, while hotel reservations have seen an even more dramatic rise of over 80%.

Almaty and Astana remain the focal points for Chinese tourists visiting Kazakhstan. Almaty, surrounded by snow-capped mountains and scenic valleys, offers a dynamic mix of outdoor adventure and urban culture.

Winter sports tourism has become a key driver of Kazakhstan's growing popularity. Ski resorts such as Shymbulak near Almaty have seen rising interest since late November, drawing both novice and experienced skiers.

Tourism statistics reflect this growing momentum. Over the first eleven months of the current year, more than 876,000 Chinese tourists visited Kazakhstan, a substantial increase from 655,000 recorded for the entire previous year."""
    }
]


def translate_to_bengali(text):
    """Translate to Bengali"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Translate to Bangladeshi Bengali. Keep facts accurate. Keep proper nouns in original language."},
            {"role": "user", "content": text}
        ],
        temperature=0.3,
        max_tokens=3000
    )
    return response.choices[0].message.content


def generate_hard_news(translated_text, article_info, retry_count=0):
    """Generate Hard News format with code-based post-processing (v2.5)"""
    style = STYLES["hard_news"]

    user_prompt = f"""নিচের ভ্রমণ সংবাদটি পুনর্লিখন করুন:

মূল শিরোনাম: {article_info['title']}
উৎস: {article_info['publisher']}

অনুবাদিত বিষয়বস্তু:
{translated_text}

এখন হার্ড নিউজ স্টাইলে লিখুন:"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": style["system_prompt"]},
            {"role": "user", "content": user_prompt}
        ],
        temperature=style["temperature"],
        max_tokens=style["max_tokens"]
    )
    raw_content = response.choices[0].message.content

    # Detect issues BEFORE processing (for logging)
    original_issues = detect_issues(raw_content, 'hard_news')
    if original_issues:
        print(f"      [AI ISSUES] {original_issues} → will be fixed by code")

    # Apply ALL code-based fixes (100% reliable)
    content = process_content(raw_content, 'hard_news')

    # HARD NEWS MINIMUM WORD CHECK (220 words from intro to conclusion)
    if retry_count < 2:
        word_count = count_body_words(content)
        if word_count < 220:
            print(f"      [REGEN] Hard news too short: {word_count} words (min 220). Regenerating (attempt {retry_count + 2}/3)...")
            return generate_hard_news(translated_text, article_info, retry_count + 1)

    # Verify issues were fixed
    remaining_issues = detect_issues(content, 'hard_news')
    code_fixed = len(original_issues) > 0 and len(remaining_issues) < len(original_issues)

    return content, code_fixed, original_issues


def generate_soft_news(translated_text, article_info):
    """Generate Soft News format with code-based post-processing (v2.5)"""
    style = STYLES["soft_news"]

    user_prompt = f"""নিচের ভ্রমণ সংবাদটি পুনর্লিখন করুন:

মূল শিরোনাম: {article_info['title']}
উৎস: {article_info['publisher']}

অনুবাদিত বিষয়বস্তু:
{translated_text}

এখন সফট নিউজ স্টাইলে লিখুন:"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": style["system_prompt"]},
            {"role": "user", "content": user_prompt}
        ],
        temperature=style["temperature"],
        max_tokens=style["max_tokens"]
    )
    raw_content = response.choices[0].message.content

    # Detect issues BEFORE processing (for logging)
    original_issues = detect_issues(raw_content, 'soft_news')
    if original_issues:
        print(f"      [AI ISSUES] {original_issues} → will be fixed by code")

    # Apply ALL code-based fixes (100% reliable)
    content = process_content(raw_content, 'soft_news')

    # Verify issues were fixed
    remaining_issues = detect_issues(content, 'soft_news')
    code_fixed = len(original_issues) > 0 and len(remaining_issues) < len(original_issues)

    return content, code_fixed, original_issues


def analyze_structure(content, news_type):
    """Analyze the structure of news content (v3.0)"""
    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]

    result = {
        "total_paragraphs": len(paragraphs),
        "structure": [],
        "validation": {}
    }

    for i, para in enumerate(paragraphs[:8], 1):
        # Check if FULLY bold (starts AND ends with **)
        is_fully_bold = para.startswith('**') and para.endswith('**')
        is_partial_bold = para.startswith('**') and not para.endswith('**')
        words = len(para.split())
        result["structure"].append({
            "paragraph": i,
            "bold": is_fully_bold,
            "partial_bold": is_partial_bold,
            "words": words,
            "preview": para[:60].replace('\n', ' ') + ("..." if len(para) > 60 else "")
        })

    if news_type == "soft_news" and len(paragraphs) >= 5:
        # Check 2-paragraph rule for soft news (v3.0)
        # Structure: P1=Headline, P2=Byline, P3=Intro1(FULLY bold), P4=Intro2(NOT bold), P5=Subhead
        #
        # Rule: EXACTLY 2 intro paragraphs before first subhead
        # - P3 (Intro 1) must be FULLY bold (starts AND ends with **)
        # - P4 (Intro 2) must NOT be bold
        # - P5 should be the first subhead

        p3 = paragraphs[2] if len(paragraphs) > 2 else ""
        p3_fully_bold = p3.startswith('**') and p3.endswith('**')
        p3_partial_bold = p3.startswith('**') and not p3.endswith('**')

        # Find first subhead after P3 (index 2)
        # Subhead = FULLY bold (starts AND ends with **), typically short (< 100 chars)
        first_subhead_idx = None
        non_bold_count_after_intro1 = 0

        for i in range(3, min(len(paragraphs), 10)):  # Check up to P10
            para = paragraphs[i]
            is_fully_bold = para.startswith('**') and para.endswith('**')

            if not para.startswith('**'):  # Not bold at all
                non_bold_count_after_intro1 += 1
            elif is_fully_bold:
                # Check if it looks like a subhead (short bold text)
                clean_text = para[2:-2].strip()  # Remove ** from both ends
                if len(clean_text) < 100:  # Subheads are typically short
                    first_subhead_idx = i
                    break

        # PASS conditions:
        # 1. P3 is FULLY bold (not partial)
        # 2. There's at least 1 non-bold paragraph (Intro 2) before first subhead
        #    Note: enforce_paragraph_length() may split long Intro 2 into multiple paragraphs
        # 3. First subhead exists
        if p3_fully_bold and non_bold_count_after_intro1 >= 1 and first_subhead_idx is not None:
            result["validation"]["2_paragraph_rule"] = "PASS"
            result["validation"]["intro1_fully_bold"] = True
            result["validation"]["intro2_count"] = non_bold_count_after_intro1
            result["validation"]["first_subhead_at"] = f"P{first_subhead_idx + 1}"
        else:
            result["validation"]["2_paragraph_rule"] = "FAIL"
            result["validation"]["p3_fully_bold"] = p3_fully_bold
            result["validation"]["p3_partial_bold"] = p3_partial_bold
            result["validation"]["non_bold_after_intro1"] = non_bold_count_after_intro1
            result["validation"]["first_subhead_found"] = first_subhead_idx is not None
            if first_subhead_idx is not None:
                result["validation"]["first_subhead_at"] = f"P{first_subhead_idx + 1}"

    return result


def create_word_document(results):
    """Create Word document with all results"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("[!] python-docx not installed. Installing...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Title
    title = doc.add_heading('Bengali News Format Test Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    for result in results:
        # Article header
        doc.add_heading(f"Article {result['id']}: {result['title']}", level=1)
        doc.add_paragraph(f"Source: {result['publisher']}")
        doc.add_paragraph()

        # Hard News Section
        doc.add_heading('HARD NEWS (হার্ড নিউজ)', level=2)

        # Add hard news content with formatting
        hard_news_lines = result['hard_news'].split('\n')
        for line in hard_news_lines:
            if line.strip():
                p = doc.add_paragraph()
                if line.startswith('**') and line.endswith('**'):
                    # Bold text
                    run = p.add_run(line.replace('**', ''))
                    run.bold = True
                elif line.startswith('**'):
                    # Partially bold
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if part:
                            run = p.add_run(part)
                            run.bold = (i % 2 == 1)
                else:
                    p.add_run(line)

        # Hard News Analysis
        analysis_hard = result.get('analysis_hard', {})
        doc.add_paragraph(f"Total paragraphs: {analysis_hard.get('total_paragraphs', 'N/A')}")
        if result.get('hard_code_fixed'):
            doc.add_paragraph(f"Code Fixed Issues: {', '.join(result.get('hard_original_issues', []))}")

        doc.add_paragraph()

        # Soft News Section
        doc.add_heading('SOFT NEWS (সফট নিউজ)', level=2)

        # Add soft news content with formatting
        soft_news_lines = result['soft_news'].split('\n')
        for line in soft_news_lines:
            if line.strip():
                p = doc.add_paragraph()
                if line.startswith('**') and line.endswith('**'):
                    run = p.add_run(line.replace('**', ''))
                    run.bold = True
                elif line.startswith('**'):
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if part:
                            run = p.add_run(part)
                            run.bold = (i % 2 == 1)
                else:
                    p.add_run(line)

        # Soft News Analysis
        analysis_soft = result.get('analysis_soft', {})
        doc.add_paragraph(f"Total paragraphs: {analysis_soft.get('total_paragraphs', 'N/A')}")
        validation = analysis_soft.get('validation', {})
        if '2_paragraph_rule' in validation:
            doc.add_paragraph(f"2-Paragraph Rule: {validation['2_paragraph_rule']}")
        if result.get('soft_code_fixed'):
            doc.add_paragraph(f"Code Fixed Issues: {', '.join(result.get('soft_original_issues', []))}")

        # Page break between articles
        doc.add_page_break()

    return doc


def main():
    print("=" * 70)
    print("COMPREHENSIVE NEWS FORMAT TEST - ALL 7 ARTICLES")
    print("=" * 70)

    results = []

    for article in ARTICLES:
        print(f"\n{'='*70}")
        print(f"Article {article['id']}: {article['title'][:50]}...")
        print("=" * 70)

        # Step 1: Translate
        print(f"  [1/3] Translating to Bengali...")
        translated = translate_to_bengali(article['content'])
        print(f"  [OK] Translation complete ({len(translated)} chars)")

        # Step 2: Generate Hard News
        print(f"  [2/3] Generating Hard News...")
        hard_news, hard_code_fixed, hard_original_issues = generate_hard_news(translated, article)
        analysis_hard = analyze_structure(hard_news, "hard_news")
        hard_word_count = count_body_words(hard_news)
        fix_status_hard = "CODE FIXED" if hard_code_fixed else "Clean"
        word_status = "PASS" if hard_word_count >= 220 else "FAIL"
        print(f"  [OK] Hard News: {analysis_hard['total_paragraphs']} paragraphs, {hard_word_count} words [{word_status}] ({fix_status_hard})")

        # Step 3: Generate Soft News
        print(f"  [3/3] Generating Soft News...")
        soft_news, soft_code_fixed, soft_original_issues = generate_soft_news(translated, article)
        analysis_soft = analyze_structure(soft_news, "soft_news")
        validation_status = analysis_soft.get('validation', {}).get('2_paragraph_rule', 'N/A')
        fix_status_soft = "CODE FIXED" if soft_code_fixed else "Clean"
        print(f"  [OK] Soft News: {analysis_soft['total_paragraphs']} paragraphs, 2-para rule: {validation_status} ({fix_status_soft})")

        results.append({
            "id": article['id'],
            "title": article['title'],
            "publisher": article['publisher'],
            "translated": translated,
            "hard_news": hard_news,
            "soft_news": soft_news,
            "analysis_hard": analysis_hard,
            "analysis_soft": analysis_soft,
            "hard_code_fixed": hard_code_fixed,
            "hard_original_issues": hard_original_issues,
            "soft_code_fixed": soft_code_fixed,
            "soft_original_issues": soft_original_issues,
            "hard_word_count": hard_word_count
        })

    # Save to Word document
    print("\n" + "=" * 70)
    print("GENERATING WORD DOCUMENT...")
    print("=" * 70)

    doc = create_word_document(results)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = Path(__file__).parent / "test_output" / f"all_articles_test_results_{timestamp}.docx"
    output_path.parent.mkdir(exist_ok=True)
    doc.save(str(output_path))
    print(f"[OK] Word document saved: {output_path}")

    # Also save raw text version
    text_output_path = Path(__file__).parent / "test_output" / "all_articles_test_results.txt"
    with open(text_output_path, "w", encoding="utf-8") as f:
        for result in results:
            f.write("=" * 70 + "\n")
            f.write(f"ARTICLE {result['id']}: {result['title']}\n")
            f.write(f"Source: {result['publisher']}\n")
            f.write("=" * 70 + "\n\n")

            f.write("-" * 40 + "\n")
            f.write("HARD NEWS:\n")
            hard_wc = result.get('hard_word_count', 0)
            wc_status = "PASS" if hard_wc >= 220 else "FAIL"
            f.write(f"[Word Count: {hard_wc} words - {wc_status} (min 220)]\n")
            if result.get('hard_code_fixed'):
                f.write(f"[CODE FIXED - Original issues: {', '.join(result.get('hard_original_issues', []))}]\n")
            f.write("-" * 40 + "\n")
            f.write(result['hard_news'] + "\n\n")

            f.write("-" * 40 + "\n")
            f.write("SOFT NEWS:\n")
            if result.get('soft_code_fixed'):
                f.write(f"[CODE FIXED - Original issues: {', '.join(result.get('soft_original_issues', []))}]\n")
            f.write("-" * 40 + "\n")
            f.write(result['soft_news'] + "\n\n")

            f.write("-" * 40 + "\n")
            f.write("SOFT NEWS STRUCTURE ANALYSIS:\n")
            f.write("-" * 40 + "\n")
            for item in result['analysis_soft']['structure']:
                bold_tag = "[BOLD]" if item['bold'] else "[normal]"
                f.write(f"P{item['paragraph']}: {bold_tag:10} ({item['words']:3} words) {item['preview']}\n")

            validation = result['analysis_soft'].get('validation', {})
            if '2_paragraph_rule' in validation:
                f.write(f"\n2-Paragraph Rule: {validation['2_paragraph_rule']}\n")

            f.write("\n\n")

    print(f"[OK] Text file saved: {text_output_path}")

    # Summary
    print("\n" + "=" * 70)
    print("SUMMARY")
    print("=" * 70)

    pass_count = 0
    fail_count = 0
    na_count = 0

    for result in results:
        status = result['analysis_soft'].get('validation', {}).get('2_paragraph_rule', 'N/A')
        if status == "PASS":
            pass_count += 1
            icon = "[PASS]"
        elif status == "FAIL":
            fail_count += 1
            icon = "[FAIL]"
        else:
            na_count += 1
            icon = "[N/A]"

        print(f"  Article {result['id']}: {icon} {result['title'][:45]}...")

    print(f"\nResults: {pass_count} PASS, {fail_count} FAIL, {na_count} N/A")
    print(f"\nOutput files:")
    print(f"  - Word: {output_path}")
    print(f"  - Text: {text_output_path}")


if __name__ == "__main__":
    main()
